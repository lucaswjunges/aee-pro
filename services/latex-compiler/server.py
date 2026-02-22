import os
import re
import base64
import subprocess
import tempfile
import shutil
import json as json_lib
import urllib.request

from fastapi import FastAPI, HTTPException, Header, BackgroundTasks, Response
from pydantic import BaseModel

app = FastAPI(title="AEE+ PRO LaTeX Compiler")

AUTH_TOKEN = os.environ.get("COMPILER_AUTH_TOKEN", "")


class ImagePayload(BaseModel):
    filename: str
    data_base64: str


class CompileRequest(BaseModel):
    latex_source: str
    images: list[ImagePayload] | None = None


class CompileResponse(BaseModel):
    success: bool
    pdf_base64: str | None = None
    pdf_size_bytes: int | None = None
    error: str | None = None
    warnings: list[str] | None = None


_WARNING_PATTERNS = [
    re.compile(r"^(Overfull \\[hv]box .+)$", re.MULTILINE),
    re.compile(r"^(Underfull \\[hv]box .+)$", re.MULTILINE),
    re.compile(r"^(LaTeX Warning: .+)$", re.MULTILINE),
    re.compile(r"^(Package \S+ Warning: .+)$", re.MULTILINE),
]

MAX_WARNINGS = 30


def _extract_warnings(log_path: str) -> list[str]:
    """Parse a LaTeX .log file and return relevant warning lines."""
    if not os.path.exists(log_path):
        return []
    with open(log_path, "r", encoding="utf-8", errors="replace") as f:
        log_text = f.read()
    warnings: list[str] = []
    for pat in _WARNING_PATTERNS:
        for m in pat.finditer(log_text):
            warnings.append(m.group(1).strip())
            if len(warnings) >= MAX_WARNINGS:
                return warnings
    return warnings


MAX_IMAGES_TOTAL_BYTES = 10 * 1024 * 1024  # 10 MB


def _prepare_images(images: list[ImagePayload] | None, tmpdir: str) -> bool:
    """Decode images to tmpdir/images/. Returns True if images were written."""
    if not images:
        return False
    images_dir = os.path.join(tmpdir, "images")
    os.makedirs(images_dir, exist_ok=True)
    total_bytes = 0
    for img in images:
        data = base64.b64decode(img.data_base64)
        total_bytes += len(data)
        if total_bytes > MAX_IMAGES_TOTAL_BYTES:
            raise ValueError(f"Total de imagens excede {MAX_IMAGES_TOTAL_BYTES // (1024*1024)}MB")
        # Sanitize filename — only allow alphanumeric, dash, underscore, dot
        safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", img.filename)
        with open(os.path.join(images_dir, safe_name), "wb") as f:
            f.write(data)
    return True


def _enable_real_graphicx(latex_source: str) -> str:
    """Replace draft graphicx with real graphicx and add graphicspath."""
    # Use regex to handle optional comments/whitespace after the command
    result = re.sub(
        r"\\usepackage\[draft\]\{graphicx\}[^\n]*",
        r"\\usepackage{graphicx}\n\\graphicspath{{./images/}}",
        latex_source,
    )
    return result


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/compile", response_model=CompileResponse)
def compile_latex(
    req: CompileRequest,
    authorization: str = Header(default=""),
):
    # Auth check
    if AUTH_TOKEN:
        token = authorization.removeprefix("Bearer ").strip()
        if token != AUTH_TOKEN:
            raise HTTPException(status_code=401, detail="Unauthorized")

    tmpdir = tempfile.mkdtemp(prefix="latex_")
    tex_path = os.path.join(tmpdir, "document.tex")
    pdf_path = os.path.join(tmpdir, "document.pdf")

    try:
        latex_source = req.latex_source

        # Decode images and enable real graphicx if images provided
        try:
            has_images = _prepare_images(req.images, tmpdir)
        except ValueError as e:
            return CompileResponse(success=False, error=str(e))
        if has_images:
            latex_source = _enable_real_graphicx(latex_source)

        # Write .tex file
        with open(tex_path, "w", encoding="utf-8") as f:
            f.write(latex_source)

        # Run pdflatex twice (for table of contents / references)
        for pass_num in range(2):
            result = subprocess.run(
                [
                    "pdflatex",
                    "-interaction=nonstopmode",
                    "-halt-on-error",
                    "-output-directory", tmpdir,
                    tex_path,
                ],
                capture_output=True,
                timeout=60,
                cwd=tmpdir,
            )

            # Decode stdout/stderr safely
            stdout = result.stdout.decode("utf-8", errors="replace") if result.stdout else ""
            stderr = result.stderr.decode("utf-8", errors="replace") if result.stderr else ""

            if result.returncode != 0:
                # Extract meaningful error lines from log
                log_path = os.path.join(tmpdir, "document.log")
                error_log = ""
                if os.path.exists(log_path):
                    with open(log_path, "r", encoding="utf-8", errors="replace") as f:
                        lines = f.readlines()
                    # Extract error lines
                    error_lines = []
                    capture = False
                    for line in lines:
                        if line.startswith("!") or capture:
                            error_lines.append(line.rstrip())
                            capture = True
                            if len(error_lines) > 5:
                                capture = False
                        if len(error_lines) > 30:
                            break
                    error_log = "\n".join(error_lines) if error_lines else stdout[-2000:]
                else:
                    error_log = stdout[-2000:] if stdout else stderr[-2000:]

                return CompileResponse(
                    success=False,
                    error=error_log[:3000],
                )

        # Check PDF exists
        if not os.path.exists(pdf_path):
            return CompileResponse(
                success=False,
                error="PDF was not generated (file not found after compilation)",
            )

        # Extract warnings from log
        log_path = os.path.join(tmpdir, "document.log")
        warnings = _extract_warnings(log_path) or None

        # Read PDF and encode as base64
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        return CompileResponse(
            success=True,
            pdf_base64=base64.b64encode(pdf_bytes).decode("ascii"),
            pdf_size_bytes=len(pdf_bytes),
            warnings=warnings,
        )

    except subprocess.TimeoutExpired:
        return CompileResponse(
            success=False,
            error="Compilation timed out (60s limit)",
        )
    except Exception as e:
        return CompileResponse(
            success=False,
            error=f"Server error: {str(e)}",
        )
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


class ConvertDocxResponse(BaseModel):
    success: bool
    docx_base64: str | None = None
    docx_size_bytes: int | None = None
    error: str | None = None


_UNTITLED_ENVS = ["datacard", "materialbox"]


def _extract_brace_arg(text: str, start: int) -> tuple[str, int]:
    """Extract content from a {...} group, handling nested braces."""
    if start >= len(text) or text[start] != "{":
        return ("", start)
    depth = 0
    i = start
    while i < len(text):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                return (text[start + 1 : i], i + 1)
        i += 1
    return (text[start + 1 :], len(text))


def _extract_cover_info(tikz_block: str) -> str:
    """Extract text content from a TikZ cover page into plain LaTeX."""
    lines: list[str] = []
    # Extract \textbf{Key:} Value patterns from \node content
    for m in re.finditer(r"\\textbf\{([^}]*)\}\s*([^\\}\n]+)", tikz_block):
        key = m.group(1).strip()
        val = m.group(2).strip()
        if key and val:
            lines.append(f"\\textbf{{{key}}} {val}")
    if not lines:
        return ""
    return "\\begin{center}\n" + " \\\\\n".join(lines) + "\n\\end{center}\n"


def _remove_adjustbox_cmd(text: str) -> str:
    """Remove \\adjustbox{options}{content} command form, keeping content."""
    while True:
        m = re.search(r"\\adjustbox\{[^}]*\}\s*\{", text)
        if not m:
            break
        # Find the opening { of the content
        brace_start = text.index("{", m.start() + len("\\adjustbox{"))
        # Skip past the options {..}
        _, after_opts = _extract_brace_arg(text, m.start() + len("\\adjustbox"))
        # Now extract the content {..}
        if after_opts < len(text) and text[after_opts] == "{":
            content, end_pos = _extract_brace_arg(text, after_opts)
            text = text[:m.start()] + content + text[end_pos:]
        else:
            break
    return text


def _extract_doc_title_from_preamble(source: str) -> tuple[str, str]:
    """Extract document title and student name from fancyhead in preamble."""
    title = ""
    student = ""
    begin_idx = source.find(r"\begin{document}")
    preamble = source[:begin_idx] if begin_idx != -1 else ""
    # \fancyhead[L]{\small\color{textgray}\textit{Sugestão de Atendimento}}
    # Use .*? with DOTALL-safe approach — match \textit inside fancyhead[L] line
    for line in preamble.split("\n"):
        if r"\fancyhead[L]" in line:
            m = re.search(r"\\textit\{([^}]+)\}", line)
            if m:
                title = m.group(1).strip()
        elif r"\fancyhead[R]" in line:
            m = re.search(r"\\textit\{([^}]+)\}", line)
            if m:
                student = m.group(1).strip()
                if "---" in student:
                    student = student.split("---")[0].strip()
    return title, student


def _preprocess_latex_for_pandoc(source: str) -> str:
    """Convert custom LaTeX to standard LaTeX that pandoc understands well."""

    # 0) Extract title/student from preamble before stripping it
    doc_title, doc_student = _extract_doc_title_from_preamble(source)

    # 1) Strip everything before \begin{document}
    begin_idx = source.find(r"\begin{document}")
    if begin_idx != -1:
        body = source[begin_idx:]
    else:
        body = source

    # 2) Handle atividadebox specially: \begin{atividadebox}[color]{★ Title}
    #    The TITLE is in the {braces}, the [brackets] is the color
    def _replace_atividadebox(m):
        rest = body[m.end():]
        # skip optional [color]
        pos = 0
        if rest and rest[0] == "[":
            close = rest.find("]")
            if close != -1:
                pos = close + 1
        # extract {Title}
        if pos < len(rest) and rest[pos] == "{":
            title, end_pos = _extract_brace_arg(rest, pos)
            # Clean up commands in title
            title = re.sub(r"\\starmark\s*~?\s*", "★ ", title)
            title = re.sub(r"\\[a-zA-Z]+\s*", "", title)
            title = title.strip()
            # Return subsection + remainder
            return f"\n\\subsection*{{{title}}}\n" + rest[end_pos:]
        return "\n" + rest

    # Process atividadebox iteratively
    while True:
        m = re.search(r"\\begin\{atividadebox\}", body)
        if not m:
            break
        before = body[: m.start()]
        replaced = _replace_atividadebox(m)
        body = before + replaced
    body = body.replace("\\end{atividadebox}", "\n")

    # 3) Handle other titled envs: \begin{env}[Title]
    #    infobox/sessaobox → \section (H1) — these are major document sections
    #    alertbox/successbox/dicabox → \subsection (H2) — these are minor callouts
    _H1_ENVS = {"sessaobox", "infobox"}
    for env in ["sessaobox", "infobox", "alertbox", "successbox", "dicabox"]:
        level = "section" if env in _H1_ENVS else "subsection"
        pattern = re.compile(
            r"\\begin\{" + env + r"\}\[([^\]]*)\]"
        )
        body = pattern.sub(lambda m, l=level: f"\n\\{l}*{{{m.group(1)}}}\n", body)
        # Also handle without title
        body = re.sub(r"\\begin\{" + env + r"\}", "\n", body)
        body = body.replace(f"\\end{{{env}}}", "\n")

    # 4) Remove untitled box environments — keep content
    for env in _UNTITLED_ENVS:
        body = re.sub(r"\\begin\{" + env + r"\}(?:\[[^\]]*\])?(?:\{[^}]*\})?", "\n", body)
        body = body.replace(f"\\end{{{env}}}", "\n")

    # 5) Remove any remaining tcolorbox environments
    body = re.sub(r"\\begin\{tcolorbox\}(?:\[[^\]]*\])?", "\n", body)
    body = body.replace("\\end{tcolorbox}", "\n")

    # 6) TikZ pictures — extract cover info from first, replace all with note
    tikz_blocks = list(re.finditer(
        r"\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}",
        body,
        flags=re.DOTALL,
    ))
    # Process in reverse order to preserve indices
    for i, m in enumerate(reversed(tikz_blocks)):
        idx = len(tikz_blocks) - 1 - i
        if idx == 0:
            # First tikzpicture is the cover — extract info
            cover_info = _extract_cover_info(m.group())
            if cover_info:
                body = body[:m.start()] + cover_info + body[m.end():]
            else:
                body = body[:m.start()] + "\n\\emph{[Diagrama visual -- ver PDF]}\n" + body[m.end():]
        else:
            body = body[:m.start()] + "\n\\emph{[Diagrama visual -- ver PDF]}\n" + body[m.end():]

    # 7) Remove adjustbox — both environment and command forms
    body = re.sub(r"\\begin\{adjustbox\}\{[^}]*\}", "", body)
    body = body.replace("\\end{adjustbox}", "")
    body = _remove_adjustbox_cmd(body)

    # 8) Replace custom icon commands with Unicode
    body = re.sub(r"\\cmark\b", "✓", body)
    body = re.sub(r"\\starmark\b\s*~?\s*", "★ ", body)
    body = re.sub(r"\\hand\b", "☞", body)
    body = re.sub(r"\\bulb\b\s*~?\s*", "💡 ", body)

    # 9) Remove \objtag — extract the text content (no brackets to avoid
    #    pandoc interpreting \item [text] as a label)
    body = re.sub(r"\\objtag(?:\[[^\]]*\])?\{([^}]*)\}", r"\1", body)

    # 10) Remove color commands — keep text
    body = re.sub(r"\\textcolor\{[^}]*\}\{([^}]*)\}", r"\1", body)
    body = re.sub(r"\\color\{[^}]*\}", "", body)
    body = re.sub(r"\\rowcolor\{[^}]*\}", "", body)
    body = re.sub(r"\\cellcolor\{[^}]*\}", "", body)
    body = re.sub(r"\\columncolor\{[^}]*\}", "", body)

    # 11) Remove \makecell — keep text
    body = re.sub(r"\\makecell(?:\[[^\]]*\])?\{([^}]*)\}", r"\1", body)

    # 12) Simplify tabularx → tabular with simple column spec
    #     Column specs can span multiple lines with nested braces
    def _replace_tabularx(m):
        rest = body[m.end():]
        # Skip {width}
        if rest.startswith("{"):
            _, pos = _extract_brace_arg(rest, 0)
        else:
            pos = 0
        # Skip {col spec} - may contain nested braces
        if pos < len(rest) and rest[pos] == "{":
            col_spec, end_pos = _extract_brace_arg(rest, pos)
            # Count columns
            n_cols = len(re.findall(r"[XlcrpL]", col_spec, re.IGNORECASE))
            if n_cols == 0:
                n_cols = col_spec.count("&") + 2  # rough guess
            n_cols = max(n_cols, 2)
            simple_spec = "|".join(["l"] * n_cols)
            return (f"\\begin{{tabular}}{{|{simple_spec}|}}", m.start(), m.end() + end_pos)
        return (f"\\begin{{tabular}}{{|l|l|}}", m.start(), m.end())

    # Process tabularx iteratively
    while True:
        m = re.search(r"\\begin\{tabularx\}", body)
        if not m:
            break
        result_tuple = _replace_tabularx(m)
        replacement, start, end = result_tuple
        body = body[:start] + replacement + body[end:]
    body = body.replace("\\end{tabularx}", "\\end{tabular}")

    # Also simplify complex tabular specs (with >{} modifiers)
    def _simplify_tabular_spec(m):
        spec = m.group(1)
        if ">" not in spec and "p{" not in spec:
            return m.group(0)  # already simple
        n_cols = spec.count("&") + 1
        if n_cols <= 1:
            n_cols = len(re.findall(r"[lcrpX]", spec, re.IGNORECASE))
        n_cols = max(n_cols, 2)
        simple = "|".join(["l"] * n_cols)
        return f"\\begin{{tabular}}{{|{simple}|}}"
    body = re.sub(r"\\begin\{tabular\}\{([^}]*)\}", _simplify_tabular_spec, body)

    # Remove \hline duplicates (pandoc handles single \hline fine)
    body = re.sub(r"(\\hline\s*){2,}", r"\\hline\n", body)

    # 13) Remove longtable-specific commands
    body = re.sub(r"\\endhead\b", "", body)
    body = re.sub(r"\\endfoot\b", "", body)
    body = re.sub(r"\\endfirsthead\b", "", body)
    body = re.sub(r"\\endlastfoot\b", "", body)

    # 14) Remove watermark
    body = re.sub(r"\\SetWatermark\w+\{[^}]*\}", "", body)

    # 15) \hrulefill → underscores (fill-in-the-blank lines)
    body = body.replace("\\hrulefill", "________________")

    # 16) Signature: Convert side-by-side minipages into a 2-column table
    #     Pattern: \begin{minipage}...content...\end{minipage}%\hfill%\begin{minipage}...content...\end{minipage}
    def _minipages_to_table(body_text: str) -> str:
        """Convert adjacent minipage pairs (joined by %\\hfill%) into a 2-column table."""
        mp_pair = re.compile(
            r"\\begin\{minipage\}(?:\[[^\]]*\])?\{[^}]*\}"  # \begin{minipage}[t]{0.45\textwidth}
            r"(.*?)"                                          # content 1
            r"\\end\{minipage\}"
            r"\s*%\\hfill%\s*"                                # %\hfill% glue
            r"\\begin\{minipage\}(?:\[[^\]]*\])?\{[^}]*\}"
            r"(.*?)"                                          # content 2
            r"\\end\{minipage\}",
            re.DOTALL,
        )
        def _parse_sig_lines(content: str) -> list[str]:
            """Extract signature lines from minipage content."""
            content = re.sub(r"\\centering\b", "", content)
            content = re.sub(r"\\rule\{[^}]*\}\{[^}]*\}", "________________________________", content)
            content = re.sub(r"\\\\\s*\[\d+pt\]", "\n", content)
            content = re.sub(r"\\\\", "\n", content)
            lines = [l.strip() for l in content.strip().split("\n") if l.strip()]
            return lines

        def _replace_pair(m):
            left_lines = _parse_sig_lines(m.group(1))
            right_lines = _parse_sig_lines(m.group(2))
            # Pad to same length
            max_len = max(len(left_lines), len(right_lines))
            while len(left_lines) < max_len:
                left_lines.append("")
            while len(right_lines) < max_len:
                right_lines.append("")
            # Build multi-row table — one row per line
            rows = []
            for l, r in zip(left_lines, right_lines):
                rows.append(f"{l} & {r} \\\\")
            return (
                "\n\\begin{tabular}{p{0.45\\textwidth} p{0.45\\textwidth}}\n"
                + "\n".join(rows) + "\n"
                + "\\end{tabular}\n"
            )

        return mp_pair.sub(_replace_pair, body_text)

    body = _minipages_to_table(body)

    # Handle remaining standalone minipages (not paired)
    body = re.sub(r"\\rule\{[^}]*\}\{[^}]*\}", "________________________________", body)
    body = re.sub(r"\\begin\{minipage\}(?:\[[^\]]*\])?\{[^}]*\}", "\n", body)
    body = body.replace("\\end{minipage}", "\n")
    body = re.sub(r"\\hfill\b", "    ", body)
    body = re.sub(r"\\vfill\b", "\n", body)

    # 17) Clean misc commands pandoc doesn't need
    body = re.sub(r"\\noindent\b", "", body)
    body = re.sub(r"\\centering\b", "", body)
    body = re.sub(r"\\(large|Large|LARGE|huge|Huge|small|footnotesize|scriptsize)\b", "", body)
    body = re.sub(r"\\vspace\*?\{[^}]*\}", "\n", body)
    body = re.sub(r"\\hspace\*?\{[^}]*\}", "", body)
    body = re.sub(r"\\newpage\b", "\n", body)
    body = re.sub(r"\\clearpage\b", "\n", body)
    body = re.sub(r"\\pagebreak\b", "\n", body)
    # Remove % line comments (but not \%)
    body = re.sub(r"(?<!\\)%[^\n]*", "", body)

    # 18) Fix orphan commas at start of lines (from removed \hspace before ", date")
    body = re.sub(r"^\s*,\s*", "", body, flags=re.MULTILINE)

    # 19) Clean excessive blank lines
    body = re.sub(r"\n{4,}", "\n\n\n", body)

    # 20) Build a minimal preamble that pandoc can work with
    title_block = ""
    if doc_title or doc_student:
        parts = []
        if doc_title:
            parts.append(f"\\title{{{doc_title}}}")
        if doc_student:
            parts.append(f"\\author{{{doc_student}}}")
        parts.append("\\date{}")
        title_block = "\n".join(parts) + "\n"

    preamble = (
        r"\documentclass[12pt,a4paper]{article}" + "\n"
        r"\usepackage[utf8]{inputenc}" + "\n"
        r"\usepackage[T1]{fontenc}" + "\n"
        r"\usepackage[brazil]{babel}" + "\n"
        r"\usepackage{booktabs}" + "\n"
        r"\usepackage{longtable}" + "\n"
        r"\usepackage{multirow}" + "\n"
        r"\usepackage{tabularx}" + "\n"
        r"\usepackage{array}" + "\n"
        r"\usepackage{enumitem}" + "\n"
        r"\usepackage{graphicx}" + "\n"
        r"\usepackage{hyperref}" + "\n"
        + title_block
    )

    # Inject \maketitle right after \begin{document} if we have title
    if doc_title or doc_student:
        body = body.replace(r"\begin{document}", r"\begin{document}" + "\n\\maketitle\n", 1)

    return preamble + "\n" + body


def _postprocess_docx(docx_path: str) -> None:
    """Apply professional AEE+ styling to the DOCX generated by pandoc."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document(docx_path)

    # Brand colors
    AEE_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
    AEE_BLUE_LIGHT = RGBColor(0x2C, 0x5F, 0x8A)
    AEE_GOLD = RGBColor(0xC9, 0xA8, 0x4C)
    TEXT_DARK = RGBColor(0x33, 0x33, 0x33)

    # --- Style headings ---
    for p in doc.paragraphs:
        if not p.style:
            continue
        sn = p.style.name

        if sn == "Heading 1":
            for run in p.runs:
                run.font.color.rgb = AEE_BLUE
                run.font.size = Pt(16)
                run.bold = True
            # Add bottom border (gold line)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{AEE_GOLD}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

        elif sn == "Heading 2":
            for run in p.runs:
                run.font.color.rgb = AEE_BLUE_LIGHT
                run.font.size = Pt(13)
                run.bold = True

        elif sn in ("Body Text", "First Paragraph", "Normal"):
            for run in p.runs:
                if not run.bold and not run.italic:
                    run.font.color.rgb = TEXT_DARK
                run.font.size = run.font.size or Pt(11)

    # --- Style tables ---
    HEADER_BG = "1E3A5F"
    ROW_ALT_1 = "F0F4F8"
    ROW_ALT_2 = "FFFFFF"
    GOLD_BG = "FFF8E1"

    for table in doc.tables:
        # Set table width to full page
        table.autofit = True

        # Skip signature tables — detect by first row being all underscores/dashes
        if table.rows:
            first_row_text = " ".join(c.text.strip() for c in table.rows[0].cells)
            if first_row_text and all(ch in "_ \t\n-" for ch in first_row_text):
                continue

        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                # Header row
                if r_idx == 0:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG}" w:val="clear"/>'
                    )
                    cell._element.get_or_add_tcPr().append(shading)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                            run.font.size = Pt(10)
                else:
                    # Alternating row colors
                    bg = GOLD_BG if r_idx % 2 == 1 else ROW_ALT_1
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>'
                    )
                    cell._element.get_or_add_tcPr().append(shading)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)
                            run.font.color.rgb = TEXT_DARK

    # --- Add header/footer ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        # Add page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")}> PAGE </w:instrText>')
        run._element.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._element.append(fldChar2)

    doc.save(docx_path)


@app.post("/convert-docx", response_model=ConvertDocxResponse)
def convert_to_docx(
    req: CompileRequest,
    authorization: str = Header(default=""),
):
    """Preprocess LaTeX (strip custom envs) then convert to DOCX via pandoc."""
    if AUTH_TOKEN:
        token = authorization.removeprefix("Bearer ").strip()
        if token != AUTH_TOKEN:
            raise HTTPException(status_code=401, detail="Unauthorized")

    tmpdir = tempfile.mkdtemp(prefix="docx_")
    tex_path = os.path.join(tmpdir, "document.tex")
    docx_path = os.path.join(tmpdir, "document.docx")

    try:
        latex_source = req.latex_source

        # Decode images for pandoc conversion
        try:
            has_images = _prepare_images(req.images, tmpdir)
        except ValueError as e:
            return ConvertDocxResponse(success=False, error=str(e))
        if has_images:
            latex_source = _enable_real_graphicx(latex_source)

        # Preprocess: convert custom LaTeX to standard LaTeX
        clean_latex = _preprocess_latex_for_pandoc(latex_source)

        with open(tex_path, "w", encoding="utf-8") as f:
            f.write(clean_latex)

        result = subprocess.run(
            [
                "pandoc",
                tex_path,
                "-f", "latex",
                "-t", "docx",
                "-o", docx_path,
                "--wrap=preserve",
            ],
            capture_output=True,
            timeout=60,
            cwd=tmpdir,
        )

        if result.returncode != 0:
            stderr = result.stderr.decode("utf-8", errors="replace") if result.stderr else ""
            return ConvertDocxResponse(
                success=False,
                error=f"Pandoc error: {stderr[:2000]}",
            )

        if not os.path.exists(docx_path):
            return ConvertDocxResponse(
                success=False,
                error="DOCX was not generated",
            )

        # Post-process: apply AEE+ PRO styling
        _postprocess_docx(docx_path)

        with open(docx_path, "rb") as f:
            docx_bytes = f.read()

        return ConvertDocxResponse(
            success=True,
            docx_base64=base64.b64encode(docx_bytes).decode("ascii"),
            docx_size_bytes=len(docx_bytes),
        )

    except subprocess.TimeoutExpired:
        return ConvertDocxResponse(
            success=False,
            error="Conversion timed out (60s limit)",
        )
    except Exception as e:
        return ConvertDocxResponse(
            success=False,
            error=f"Server error: {str(e)}",
        )
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


# ---------------------------------------------------------------------------
# POST /generate-and-compile — Claude API + pdflatex in one shot
# ---------------------------------------------------------------------------

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL = os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6")

import logging
_log = logging.getLogger("aee-latex")


class GenerateAndCompileRequest(BaseModel):
    system_prompt: str
    user_prompt: str
    preamble: str
    max_tokens: int = 16000
    images: list[ImagePayload] | None = None
    signature_block: str | None = None
    # Webhook fields — if set, endpoint returns 202 and processes asynchronously
    doc_id: str = ""
    callback_url: str = ""
    callback_token: str = ""
    # Fallback credentials — used if service ANTHROPIC_API_KEY is out of credits
    fallback_api_key: str = ""
    fallback_model: str = ""


class GenerateAndCompileResponse(BaseModel):
    success: bool
    pdf_base64: str | None = None
    pdf_size_bytes: int | None = None
    latex_source: str | None = None
    error: str | None = None
    warnings: list[str] | None = None
    attempts: int = 0
    ai_model: str | None = None


def _extract_latex_body(raw: str) -> str:
    """Extract body from AI response (\\begin{document} to \\end{document})."""
    cleaned = re.sub(r"```latex\s*", "", raw)
    cleaned = re.sub(r"```\s*", "", cleaned).strip()

    start = cleaned.find("\\begin{document}")
    end = cleaned.rfind("\\end{document}")

    if start != -1 and end != -1 and end > start:
        body = cleaned[start : end + len("\\end{document}")]
    elif start != -1:
        body = cleaned[start:] + "\n\\end{document}"
    else:
        body = "\\begin{document}\n" + cleaned
        if not body.rstrip().endswith("\\end{document}"):
            body += "\n\\end{document}"

    return body


def _sanitize_latex(source: str) -> str:
    """Remove problematic LaTeX constructs that commonly break compilation."""
    # Strip emoji and supplemental Unicode characters that pdflatex cannot process.
    # pdflatex only supports characters declared in utf8.def (Latin-1 + some extensions).
    # Emoji (U+1F000+), Misc Symbols (U+2600-U+27BF), and other high-plane characters
    # cause a fatal "Unicode character not set up for use with LaTeX" error.
    source = re.sub(
        r"[\U00002600-\U000027BF\U0001F000-\U0010FFFF]",
        "", source,
    )
    # Remove \\foreach blocks with rnd
    source = re.sub(
        r"\\foreach\s+\\[a-zA-Z]+\s+in\s*\{[^}]*rnd[^}]*\}[^\n]*\n?",
        "", source,
    )
    # Remove TeX conditionals
    source = re.sub(r"\\if(?:num|dim|x|odd|case)\b[^\n]*\n?", "", source)
    source = re.sub(r"^\s*\\(?:else|or)\s*$", "", source, flags=re.MULTILINE)
    source = re.sub(r"^\s*\\fi\b\s*$", "", source, flags=re.MULTILINE)
    # Remove inline \\pgfmathparse in color specs
    source = re.sub(r"\\pgfmathparse\{[^}]*\}\\pgfmathresult", "50", source)
    # Close unclosed environments
    env_regex = re.compile(r"\\(begin|end)\{([^}]+)\}")
    stack: list[str] = []
    for m in env_regex.finditer(source):
        if m.group(1) == "begin":
            stack.append(m.group(2))
        elif m.group(1) == "end" and stack and stack[-1] == m.group(2):
            stack.pop()
    if stack:
        end_doc_idx = source.rfind("\\end{document}")
        insert_point = end_doc_idx if end_doc_idx != -1 else len(source)
        closings = "\n".join(f"\\end{{{env}}}" for env in reversed(stack) if env != "document")
        if closings:
            source = source[:insert_point] + "\n" + closings + "\n" + source[insert_point:]
    return source


def _compile_in_tmpdir(
    latex_source: str,
    images: list[ImagePayload] | None = None,
) -> CompileResponse:
    """Compile LaTeX in a fresh tmpdir. Handles cleanup."""
    tmpdir = tempfile.mkdtemp(prefix="gencomp_")
    try:
        tex_path = os.path.join(tmpdir, "document.tex")
        pdf_path = os.path.join(tmpdir, "document.pdf")

        source = latex_source
        try:
            has_images = _prepare_images(images, tmpdir)
        except ValueError as e:
            return CompileResponse(success=False, error=str(e))
        if has_images:
            source = _enable_real_graphicx(source)

        with open(tex_path, "w", encoding="utf-8") as f:
            f.write(source)

        for _ in range(2):
            result = subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "-halt-on-error",
                 "-output-directory", tmpdir, tex_path],
                capture_output=True, timeout=60, cwd=tmpdir,
            )
            if result.returncode != 0:
                log_path = os.path.join(tmpdir, "document.log")
                stdout = result.stdout.decode("utf-8", errors="replace") if result.stdout else ""
                stderr = result.stderr.decode("utf-8", errors="replace") if result.stderr else ""
                error_log = ""
                if os.path.exists(log_path):
                    with open(log_path, "r", encoding="utf-8", errors="replace") as f:
                        lines = f.readlines()
                    error_lines = []
                    capture = False
                    for line in lines:
                        if line.startswith("!") or capture:
                            error_lines.append(line.rstrip())
                            capture = True
                            if len(error_lines) > 5:
                                capture = False
                        if len(error_lines) > 30:
                            break
                    error_log = "\n".join(error_lines) if error_lines else stdout[-2000:]
                else:
                    error_log = stdout[-2000:] if stdout else stderr[-2000:]
                return CompileResponse(success=False, error=error_log[:3000])

        if not os.path.exists(pdf_path):
            return CompileResponse(success=False, error="PDF not generated")

        log_path = os.path.join(tmpdir, "document.log")
        warnings = _extract_warnings(log_path) or None
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        return CompileResponse(
            success=True,
            pdf_base64=base64.b64encode(pdf_bytes).decode("ascii"),
            pdf_size_bytes=len(pdf_bytes),
            warnings=warnings,
        )
    except subprocess.TimeoutExpired:
        return CompileResponse(success=False, error="Compilation timed out (60s)")
    except Exception as e:
        return CompileResponse(success=False, error=f"Server error: {str(e)}")
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


_NOISE_WARNINGS = [
    "Label(s) may have changed",
    "There were multiply-defined labels",
    "destination with the same identifier",
    "Rerun to get",
]


def _filter_significant_warnings(warnings: list[str]) -> list[str]:
    """Keep only warnings worth fixing; drop noise-only entries."""
    return [w for w in warnings if not any(n in w for n in _NOISE_WARNINGS)]


AUTOFIX_WARNINGS_SYSTEM = """Você é um especialista em LaTeX. O código abaixo compilou com sucesso, mas gerou os avisos listados.

REGRAS — corrija os avisos SEM alterar o conteúdo do documento:
1. Overfull \\hbox: reduza largura de colunas ou envolva tabular em \\adjustbox{max width=\\linewidth}{...}.
2. Overfull \\vbox (tabela longa): converta \\begin{tabular} para \\begin{longtable} adicionando \\endfirsthead/\\endhead.
3. Font shape undefined (T1/cmss/b/n, TS1/...): substitua caracteres ² por \\textsuperscript{2}, ³ por \\textsuperscript{3}.
4. enumitem "Negative labelwidth": troque leftmargin=0pt por leftmargin=1.5em em description.
5. Underfull \\hbox: aumente larguras de colunas estreitas ou adicione \\raggedright na coluna.
6. LaTeX Font Warning "not available ... substituted": adicione \\normalfont ou troque \\bfseries por \\mdseries naquele contexto.
7. NUNCA altere o conteúdo textual (seções, itens, dados).
8. NUNCA coloque longtable dentro de adjustbox, tcolorbox ou qualquer grupo.

Retorne o código LaTeX COMPLETO corrigido (de \\begin{document} até \\end{document}), sem explicações, sem fence blocks."""


AUTOFIX_SYSTEM = """Você é um especialista em LaTeX. O código abaixo falhou na compilação com pdflatex.

REGRAS DE CORREÇÃO:
1. Corrija os erros mantendo o conteúdo e estilo.
2. Se o documento está TRUNCADO (texto cortado), COMPLETE o conteúdo faltante.
3. NUNCA use \\begin{axis}/pgfplots — substitua por tabelas com booktabs.
4. NUNCA coloque longtable dentro de adjustbox, tcolorbox ou qualquer grupo.
5. NUNCA use condicionais TeX (\\ifnum, \\ifcase, \\else, \\fi, \\or).
6. NUNCA use \\foreach com rnd ou \\pgfmathparse inline em cores.
7. NUNCA use \\multirowcell — use \\multirow{N}{*}{texto}.
8. Todas as tcolorbox (infobox, alertbox, etc.) já são breakable — NÃO adicione breakable.
9. \\rowcolor DEVE ser o PRIMEIRO comando de uma linha de tabela (nunca após &).
10. Para tabelas que transbordam: envolva tabular em \\adjustbox{max width=\\linewidth}{...}.
11. Se usar tabularx, SEMPRE inclua pelo menos uma coluna X.
12. NUNCA use colunas X em longtable — X é exclusivo de tabularx.

Retorne o código LaTeX corrigido COMPLETO (de \\begin{document} até \\end{document}), sem explicações, sem fence blocks."""


def _is_credit_error(e: Exception) -> bool:
    """Return True if the exception is an Anthropic credit exhaustion error."""
    return "credit balance is too low" in str(e).lower()


def _do_generate_and_compile(req: GenerateAndCompileRequest) -> dict:
    """Sync: generate LaTeX with Claude, compile + auto-fix. Returns dict."""
    import anthropic

    # --- Step 1: Generate LaTeX with Claude (streaming) ---
    # Try primary service key first; fall back to user key on credit error.
    keys_to_try: list[tuple[str, str]] = [(ANTHROPIC_API_KEY, CLAUDE_MODEL)]
    if req.fallback_api_key:
        fallback_model = req.fallback_model or CLAUDE_MODEL
        keys_to_try.append((req.fallback_api_key, fallback_model))

    ai_content = None
    ai_model = CLAUDE_MODEL
    last_gen_error = None

    for api_key, model in keys_to_try:
        if not api_key:
            continue
        client = anthropic.Anthropic(api_key=api_key)
        _log.info(f"[generate] doc_id={req.doc_id!r} Calling {model} (max_tokens={req.max_tokens}, key=...{api_key[-6:]})")
        try:
            with client.messages.stream(
                model=model,
                max_tokens=req.max_tokens,
                temperature=0.7,
                system=req.system_prompt,
                messages=[{"role": "user", "content": req.user_prompt}],
            ) as stream:
                ai_content = stream.get_final_text()
                ai_model = stream.get_final_message().model
            _log.info(f"[generate] Claude returned {len(ai_content)} chars")
            last_gen_error = None
            break  # success — stop trying keys
        except Exception as e:
            last_gen_error = e
            if _is_credit_error(e) and req.fallback_api_key and api_key != req.fallback_api_key:
                _log.warning(f"[generate] doc_id={req.doc_id!r} Credit exhausted on service key — retrying with user key")
                continue
            _log.error(f"[generate] Claude API error: {e}")
            break

    if ai_content is None:
        err_msg = f"Claude API error: {str(last_gen_error)}" if last_gen_error else "Claude API error: no key available"
        return {"success": False, "error": err_msg, "ai_model": ai_model, "attempts": 0}

    # --- Step 2: Extract body, sanitize, assemble ---
    body = _extract_latex_body(ai_content)

    if req.signature_block:
        end_doc_idx = body.rfind("\\end{document}")
        if end_doc_idx != -1:
            vfill_idx = body.rfind("\\vfill", 0, end_doc_idx)
            insert_idx = vfill_idx if vfill_idx != -1 else end_doc_idx
            body = body[:insert_idx] + "\n" + req.signature_block + "\n\n" + body[insert_idx:]

    current_source = _sanitize_latex(req.preamble + body)

    # --- Step 3: Compile → Claude fixes → recompile loop (up to 5 attempts) ---
    MAX_ATTEMPTS = 5
    last_error = None

    for attempt in range(1, MAX_ATTEMPTS + 1):
        _log.info(f"[compile] doc_id={req.doc_id!r} Attempt {attempt}/{MAX_ATTEMPTS}...")
        result = _compile_in_tmpdir(current_source, req.images)

        if result.success and result.pdf_base64:
            _log.info(f"[compile] doc_id={req.doc_id!r} SUCCESS attempt {attempt}! PDF={result.pdf_size_bytes} bytes")

            # --- Step 4: Warning-fix loop (up to 2 passes) ---
            best_source = current_source
            best_pdf_b64 = result.pdf_base64
            best_pdf_size = result.pdf_size_bytes
            best_warnings = result.warnings

            significant = _filter_significant_warnings(result.warnings or [])
            MAX_WARN_FIXES = 2
            for wfix in range(1, MAX_WARN_FIXES + 1):
                if not significant:
                    break
                _log.info(f"[warn-fix] doc_id={req.doc_id!r} pass {wfix}/{MAX_WARN_FIXES}: {len(significant)} significant warning(s)")
                try:
                    with client.messages.stream(
                        model=ai_model,
                        max_tokens=req.max_tokens,
                        temperature=0.2,
                        system=AUTOFIX_WARNINGS_SYSTEM,
                        messages=[{
                            "role": "user",
                            "content": (
                                "AVISOS DE COMPILAÇÃO:\n"
                                + "\n".join(significant)
                                + "\n\nCÓDIGO LATEX:\n"
                                + best_source
                            ),
                        }],
                    ) as wfix_stream:
                        wfix_text = wfix_stream.get_final_text()

                    wfix_body = _extract_latex_body(wfix_text)
                    wfix_body = _sanitize_latex(wfix_body)
                    preamble_end = best_source.find("\\begin{document}")
                    wfix_source = (
                        best_source[:preamble_end] + wfix_body
                        if preamble_end != -1
                        else req.preamble + wfix_body
                    )
                    wfix_result = _compile_in_tmpdir(wfix_source, req.images)
                    if wfix_result.success and wfix_result.pdf_base64:
                        best_source = wfix_source
                        best_pdf_b64 = wfix_result.pdf_base64
                        best_pdf_size = wfix_result.pdf_size_bytes
                        best_warnings = wfix_result.warnings
                        significant = _filter_significant_warnings(wfix_result.warnings or [])
                        _log.info(f"[warn-fix] doc_id={req.doc_id!r} pass {wfix} OK, remaining significant: {len(significant)}")
                    else:
                        _log.warning(f"[warn-fix] doc_id={req.doc_id!r} pass {wfix} broke compilation — keeping previous version")
                        break
                except Exception as wfix_err:
                    _log.error(f"[warn-fix] doc_id={req.doc_id!r} Claude call failed: {wfix_err}")
                    break

            return {
                "success": True,
                "pdf_base64": best_pdf_b64,
                "pdf_size_bytes": best_pdf_size,
                "latex_source": best_source,
                "warnings": best_warnings,
                "attempts": attempt,
                "ai_model": ai_model,
            }

        last_error = result.error
        _log.warning(f"[compile] doc_id={req.doc_id!r} Attempt {attempt} FAILED: {(result.error or '')[:200]}")

        if attempt == MAX_ATTEMPTS:
            break

        # Ask Claude to fix the error
        _log.info(f"[auto-fix] doc_id={req.doc_id!r} Asking Claude to fix...")
        try:
            with client.messages.stream(
                model=ai_model,
                max_tokens=req.max_tokens,
                temperature=0.2,
                system=AUTOFIX_SYSTEM,
                messages=[{
                    "role": "user",
                    "content": f"ERRO DE COMPILAÇÃO:\n{result.error}\n\nCÓDIGO LATEX COM ERRO:\n{current_source}",
                }],
            ) as fix_stream:
                fix_text = fix_stream.get_final_text()
            fixed_body = _extract_latex_body(fix_text)
            fixed_body = _sanitize_latex(fixed_body)
            _log.info(f"[auto-fix] doc_id={req.doc_id!r} Claude returned fix ({len(fixed_body)} chars)")

            preamble_end = current_source.find("\\begin{document}")
            if preamble_end != -1:
                current_source = current_source[:preamble_end] + fixed_body
            else:
                current_source = req.preamble + fixed_body
        except Exception as fix_err:
            _log.error(f"[auto-fix] doc_id={req.doc_id!r} Claude call failed: {fix_err}")
            last_error = f"Auto-fix failed: {str(fix_err)}"
            break

    _log.warning(f"[generate] doc_id={req.doc_id!r} All {MAX_ATTEMPTS} attempts failed")
    return {
        "success": False,
        "latex_source": current_source,
        "error": last_error,
        "attempts": MAX_ATTEMPTS,
        "ai_model": ai_model,
    }


def _send_callback(callback_url: str, callback_token: str, result: dict) -> None:
    """POST result dict to callback_url with Bearer auth token."""
    try:
        payload = json_lib.dumps(result).encode("utf-8")
        http_req = urllib.request.Request(
            callback_url,
            data=payload,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {callback_token}",
                # Cloudflare blocks Python-urllib (error 1010 bot protection)
                "User-Agent": "AEE-Pro-Compiler/1.0",
            },
            method="POST",
        )
        urllib.request.urlopen(http_req, timeout=30)
        _log.info(f"[callback] Sent to {callback_url}")
    except Exception as e:
        _log.error(f"[callback] Failed to send to {callback_url}: {e}")


def _process_and_callback(req: GenerateAndCompileRequest) -> None:
    """Background task: generate+compile then call webhook."""
    result = _do_generate_and_compile(req)
    if req.callback_url:
        _send_callback(req.callback_url, req.callback_token, result)


@app.post("/generate-and-compile")
def generate_and_compile(
    req: GenerateAndCompileRequest,
    background_tasks: BackgroundTasks,
    authorization: str = Header(default=""),
):
    """Generate LaTeX with Claude, compile locally, auto-fix iteratively.

    If callback_url is set: returns 202 immediately and processes in background,
    POSTing the result to callback_url when done.
    Otherwise: processes synchronously and returns the result directly.
    """
    if AUTH_TOKEN:
        token = authorization.removeprefix("Bearer ").strip()
        if token != AUTH_TOKEN:
            raise HTTPException(status_code=401, detail="Unauthorized")

    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=503, detail="ANTHROPIC_API_KEY not configured")

    if req.callback_url:
        # Async mode: acknowledge immediately, process in background thread
        background_tasks.add_task(_process_and_callback, req)
        return Response(
            content=json_lib.dumps({"status": "accepted", "doc_id": req.doc_id}),
            status_code=202,
            media_type="application/json",
        )

    # Sync mode (backward compat / local dev): process and return result
    result = _do_generate_and_compile(req)
    return Response(
        content=json_lib.dumps(result),
        status_code=200,
        media_type="application/json",
    )


# ---------------------------------------------------------------------------
# POST /compile-dossie — assemble a student dossier from existing PDFs
# ---------------------------------------------------------------------------

class DossiePdfPayload(BaseModel):
    title: str
    data_base64: str


class CompileDossieRequest(BaseModel):
    student_name: str
    student_school: str | None = None
    student_diagnosis: str | None = None
    student_grade: str | None = None
    pdfs: list[DossiePdfPayload]


class CompileDossieResponse(BaseModel):
    success: bool
    pdf_base64: str | None = None
    pdf_size_bytes: int | None = None
    error: str | None = None


MAX_DOSSIE_DOCS = 30
MAX_DOSSIE_PDF_BYTES = 20 * 1024 * 1024  # 20 MB per PDF


def _escape_latex_str(s: str) -> str:
    """Escape special LaTeX characters in a plain string."""
    replacements = [
        ("\\", r"\textbackslash{}"),
        ("&", r"\&"),
        ("%", r"\%"),
        ("$", r"\$"),
        ("#", r"\#"),
        ("_", r"\_"),
        ("{", r"\{"),
        ("}", r"\}"),
        ("~", r"\textasciitilde{}"),
        ("^", r"\textasciicircum{}"),
    ]
    for old, new in replacements:
        s = s.replace(old, new)
    return s


def _build_dossie_latex(req: CompileDossieRequest, pdf_filenames: list[str]) -> str:
    """Build a LaTeX wrapper that assembles a dossier from individual PDFs."""
    student = _escape_latex_str(req.student_name)
    school = _escape_latex_str(req.student_school or "")
    diagnosis = _escape_latex_str(req.student_diagnosis or "")
    grade = _escape_latex_str(req.student_grade or "")

    # Build document list for cover page
    doc_items = ""
    for i, pdf_payload in enumerate(req.pdfs):
        title = _escape_latex_str(pdf_payload.title)
        doc_items += f"    \\item {title}\n"

    n_docs = len(req.pdfs)
    year = "2026"
    try:
        from datetime import date
        year = str(date.today().year)
    except Exception:
        pass

    # Build info lines for cover
    info_lines = []
    info_lines.append(f"\\textbf{{Aluno(a):}} {student}")
    if school:
        info_lines.append(f"\\textbf{{Escola:}} {school}")
    if grade:
        info_lines.append(f"\\textbf{{Ano/Série:}} {grade}")
    if diagnosis:
        info_lines.append(f"\\textbf{{Diagnóstico:}} {diagnosis}")
    info_lines.append(f"\\textbf{{Documentos:}} {n_docs}")

    info_nodes = ""
    y_offset = 0
    for line in info_lines:
        info_nodes += f"      \\node[anchor=west] at (1.2, {3.5 - y_offset}) {{{\\large {line}}};\n"
        y_offset += 0.8

    # ToC entries + includepdf for each document
    include_sections = ""
    for i, (pdf_payload, fname) in enumerate(zip(req.pdfs, pdf_filenames)):
        title = _escape_latex_str(pdf_payload.title)
        include_sections += f"""
%% --- Document {i+1}: {pdf_payload.title} ---
\\addcontentsline{{toc}}{{section}}{{{title}}}
\\includepdf[pages=-,pagecommand={{\\thispagestyle{{fancy}}}}]{{{fname}}}
"""

    latex = f"""\\documentclass[a4paper,12pt]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[T1]{{fontenc}}
\\usepackage[brazil]{{babel}}
\\usepackage[margin=2cm]{{geometry}}
\\usepackage{{pdfpages}}
\\usepackage{{fancyhdr}}
\\usepackage{{tocloft}}
\\usepackage{{tikz}}
\\usetikzlibrary{{positioning,calc,shadows}}
\\usepackage{{enumitem}}
\\usepackage{{hyperref}}

% Colors
\\definecolor{{aeeblue}}{{HTML}}{{1E3A5F}}
\\definecolor{{aeegold}}{{HTML}}{{C9A84C}}
\\definecolor{{aeelightblue}}{{HTML}}{{E8F0FE}}
\\definecolor{{textgray}}{{HTML}}{{333333}}

% Header/footer
\\pagestyle{{fancy}}
\\fancyhf{{}}
\\fancyhead[L]{{\\small\\color{{textgray}}\\textit{{Dossiê — {student}}}}}
\\fancyhead[R]{{\\small\\color{{textgray}}\\textit{{AEE+ PRO}}}}
\\fancyfoot[C]{{\\small\\color{{textgray}}\\thepage}}
\\renewcommand{{\\headrulewidth}}{{0.4pt}}
\\renewcommand{{\\footrulewidth}}{{0pt}}

% ToC styling
\\renewcommand{{\\cftsecfont}}{{\\color{{aeeblue}}\\bfseries}}
\\renewcommand{{\\cftsecpagefont}}{{\\color{{aeeblue}}}}
\\renewcommand{{\\cftsecleader}}{{\\cftdotfill{{\\cftdotsep}}}}

\\hypersetup{{
  colorlinks=true,
  linkcolor=aeeblue,
  urlcolor=aeeblue,
}}

\\begin{{document}}

%% ========== COVER PAGE ==========
\\thispagestyle{{empty}}
\\begin{{tikzpicture}}[remember picture, overlay]
  % Blue header band
  \\fill[aeeblue] (current page.north west) rectangle ([yshift=-4cm]current page.north east);
  % Gold accent line
  \\fill[aeegold] ([yshift=-4cm]current page.north west) rectangle ([yshift=-4.3cm]current page.north east);

  % Title on blue band
  \\node[anchor=west, white, font=\\Huge\\bfseries] at ([xshift=2cm, yshift=-2cm]current page.north west)
    {{Dossiê do Aluno}};
  \\node[anchor=west, aeegold, font=\\large] at ([xshift=2cm, yshift=-3cm]current page.north west)
    {{Atendimento Educacional Especializado}};

  % Student info card
  \\node[
    anchor=north west,
    draw=aeeblue!30,
    fill=aeelightblue,
    rounded corners=8pt,
    minimum width=14cm,
    inner sep=15pt,
    drop shadow={{shadow xshift=1pt, shadow yshift=-1pt, opacity=0.15}},
  ] at ([xshift=2cm, yshift=-6cm]current page.north west) {{
    \\begin{{minipage}}{{13cm}}
{info_nodes}
    \\end{{minipage}}
  }};

  % Document list
  \\node[
    anchor=north west,
    font=\\large\\bfseries\\color{{aeeblue}},
  ] at ([xshift=2cm, yshift=-{10 + len(info_lines) * 0.5}cm]current page.north west)
    {{Documentos incluídos:}};

  \\node[
    anchor=north west,
    text width=14cm,
  ] at ([xshift=2cm, yshift=-{11 + len(info_lines) * 0.5}cm]current page.north west) {{
    \\begin{{enumerate}}[leftmargin=1.5em, itemsep=2pt]
{doc_items}    \\end{{enumerate}}
  }};

  % Footer
  \\node[anchor=south, font=\\small\\color{{textgray}}] at ([yshift=2cm]current page.south)
    {{Gerado por AEE+ PRO — {year}}};

  % Gold bottom line
  \\fill[aeegold] ([yshift=1.2cm]current page.south west) rectangle ([yshift=1.5cm]current page.south east);
\\end{{tikzpicture}}

\\clearpage

%% ========== TABLE OF CONTENTS ==========
\\tableofcontents
\\clearpage

%% ========== INCLUDED DOCUMENTS ==========
{include_sections}

\\end{{document}}
"""
    return latex


@app.post("/compile-dossie", response_model=CompileDossieResponse)
def compile_dossie(
    req: CompileDossieRequest,
    authorization: str = Header(default=""),
):
    """Assemble multiple PDFs into a single dossier with cover page and ToC."""
    if AUTH_TOKEN:
        token = authorization.removeprefix("Bearer ").strip()
        if token != AUTH_TOKEN:
            raise HTTPException(status_code=401, detail="Unauthorized")

    if len(req.pdfs) == 0:
        return CompileDossieResponse(success=False, error="Nenhum documento fornecido")

    if len(req.pdfs) > MAX_DOSSIE_DOCS:
        return CompileDossieResponse(
            success=False,
            error=f"Máximo de {MAX_DOSSIE_DOCS} documentos por dossiê",
        )

    tmpdir = tempfile.mkdtemp(prefix="dossie_")
    try:
        # Write each PDF to tmpdir
        pdf_filenames: list[str] = []
        for i, pdf_payload in enumerate(req.pdfs):
            data = base64.b64decode(pdf_payload.data_base64)
            if len(data) > MAX_DOSSIE_PDF_BYTES:
                return CompileDossieResponse(
                    success=False,
                    error=f"PDF '{pdf_payload.title}' excede {MAX_DOSSIE_PDF_BYTES // (1024*1024)}MB",
                )
            fname = f"doc{i:03d}.pdf"
            pdf_filenames.append(fname)
            with open(os.path.join(tmpdir, fname), "wb") as f:
                f.write(data)

        # Build LaTeX wrapper
        latex_source = _build_dossie_latex(req, pdf_filenames)
        tex_path = os.path.join(tmpdir, "dossie.tex")
        pdf_path = os.path.join(tmpdir, "dossie.pdf")

        with open(tex_path, "w", encoding="utf-8") as f:
            f.write(latex_source)

        # Compile twice (for ToC resolution)
        for pass_num in range(2):
            result = subprocess.run(
                [
                    "pdflatex",
                    "-interaction=nonstopmode",
                    "-halt-on-error",
                    "-output-directory", tmpdir,
                    tex_path,
                ],
                capture_output=True,
                timeout=120,
                cwd=tmpdir,
            )

            if result.returncode != 0:
                stdout = result.stdout.decode("utf-8", errors="replace") if result.stdout else ""
                log_path = os.path.join(tmpdir, "dossie.log")
                error_log = ""
                if os.path.exists(log_path):
                    with open(log_path, "r", encoding="utf-8", errors="replace") as f:
                        lines = f.readlines()
                    error_lines = []
                    capture = False
                    for line in lines:
                        if line.startswith("!") or capture:
                            error_lines.append(line.rstrip())
                            capture = True
                            if len(error_lines) > 5:
                                capture = False
                        if len(error_lines) > 30:
                            break
                    error_log = "\n".join(error_lines) if error_lines else stdout[-2000:]
                else:
                    error_log = stdout[-2000:]
                return CompileDossieResponse(
                    success=False,
                    error=f"Falha na compilação (pass {pass_num + 1}): {error_log[:3000]}",
                )

        if not os.path.exists(pdf_path):
            return CompileDossieResponse(
                success=False,
                error="PDF do dossiê não foi gerado",
            )

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        return CompileDossieResponse(
            success=True,
            pdf_base64=base64.b64encode(pdf_bytes).decode("ascii"),
            pdf_size_bytes=len(pdf_bytes),
        )

    except subprocess.TimeoutExpired:
        return CompileDossieResponse(
            success=False,
            error="Compilação do dossiê excedeu o tempo limite (120s)",
        )
    except Exception as e:
        return CompileDossieResponse(
            success=False,
            error=f"Erro no servidor: {str(e)}",
        )
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080)
